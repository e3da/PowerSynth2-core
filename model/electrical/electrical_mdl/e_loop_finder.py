# This is the interface where the user need to define some loop for each net_group. A graphing algorithm will be used to define loop\
#    direction for each rectangular traces.
# 1. Get initial layout generated from CS
# 2. Depending on nets on each netgroup, form a graph structure for each layer and ask user to define loop direction.
# 3. Simplify the structure into RL rects and Cap rects. 
# 4. Reuse the Loop definition every time for accelerated mesh generation
from typing import OrderedDict
import networkx as nx
import sys, os
#import dill
from networkx.readwrite import edgelist
import pandas as pd
from core.model.electrical.meshing.MeshAlgorithm import TraceIslandMesh,LayerMesh

from core.model.electrical.meshing.MeshObjects import TraceCell,RectCell
from core.model.electrical.electrical_mdl.e_loop_element import LoopEval
import numpy as np
import matplotlib.pyplot as plt
from copy import deepcopy
import pickle

from datetime import date
from multiprocessing import Pool,cpu_count
import math
import io
from enum import Enum

class StackType(Enum):
    # Types for characterization decision
    SDG = 1 # S-D-G
    GDS = 2 # G-D-S
    SDS = 3 # S-D-S
class EdgeData():
    def __init__(self):
        self.parent = None # parent trace cell
        self.dir = 1 # 1: left to right, 2: bot to top, 3: z- to z+. negative integer for the opposite dir.

class LayoutLoopInterface():
    def __init__(self,islands = None, hier_E = None, freq =1e6, layer_stack =None):
        """_summary_

        Args:
            islands (_type_, optional): _description_. Defaults to None.
            hier_E (_type_, optional): _description_. Defaults to None.
            freq (_type_, optional): _description_. Defaults to 1e6.
            layer_stack (_type_, optional): _description_. Defaults to None.
        """
        self.frequency = freq
        self.short_bw = False # Set true to short all bw for trace validation
        self.layout_info = islands
        self.layer_stack = layer_stack
        self.layer_stack_simplified = [] # simplified version which care about Signal, Ground and Dielectric only 
        self.layer_mesh_table = {}
        self.hier = hier_E
        # Layers and Netlist Handling
        self.num_RL_layer = 0 # 1 for 2D case         -- these are routing layers 
        self.num_backside_layer = 0 # 0-1 for 2D case -- these are backside layers
        self.num_C_layer = 0 # 0-1 for 2D case        -- these are dielectric layers
        
        self.graph = nx.Graph()
        self.ori_map = {}
        self.comp_nodes = {}
        self.comp_dict = {}  # Use to remember the component that has its graph built (so we dont do it again)
        self.comp_net_id = {}
        self.node_id = 0
        # all list to store x and y trace cell with defined directions:
        self.x_cells = []
        self.y_cells = []
        # all list to store x and y bondwires cell
        self.x_wires = {}
        self.y_wires = {}
        
        # networkx vars:
        self.pos = {} # to store location of the node on the graph # 2D pos only
        self.pos3d = {}
        self.nodes_dict_3d = {} # inversion of the self.pos3d above 
        self.tbl_isl_loc_net = {} # 3D table for electrical nets: isl_name -->3d pos --> node id before contraction
        self.digraph = nx.DiGraph()
        # Connect wire mesh on all trace level 
        self.tbl_isl_node_connection= {} #  check if there are full connection on all directions
        self.ribbon_dict = {} # represents bondwire group in form of ribbons and store the edge - ribbon data here
        self.tc_to_edges_init = {} # initial trace cell to edges
        self.tc_to_edges_splitted = {} # trace cell to edges duing bundle creation.

        self.mutual_pair = {} # for PEEC evaluation
        self.mutual_coup_coeff_pair = {} # for extraction into LtSpice
        self.edge_to_x_bundle = {}
        self.edge_to_y_bundle = {}
        self.contracted_nodes = {} # dict of nodes pair that can be combined
        self.contracted_edges = [] # dict of removed edges from bundle creation
        self.new_id_locs = {} # map node id to its location
        self.locs_new_id = {} # map location to its id
        self.z_list =[]

        self.doc_report = None
        self.debug = False # Turn to True to report mode. It will write all info to report.docx in the same directory. Currently working with single layout evaluation

        if self.debug:
            from docx import Document
            from docx.shared import Inches
                
    def check_number_of_electrical_layer(self):
        """
        Read the layerstack and and rearrange the layers for quick access
        """
        
        all_layer_info = self.layer_stack.all_layers_info
        for layer in all_layer_info:
            layer_obj = all_layer_info[layer]
            if layer_obj.e_type == 'S': # PowerSynth way to check for a routing layer TODO: need to specify a different way in the future
                self.num_RL_layer += 1 
                self.layer_stack_simplified.append({'id':layer_obj.id,'electrical_type':layer_obj.e_type}) # only store the ids and e_type here. 
                # if we have to characterize we know which ID to apply the model to
            elif layer_obj.e_type == 'G':
                self.num_backside_layer +=1
                self.layer_stack_simplified.append({'id':layer_obj.id,'electrical_type':layer_obj.e_type})
            elif layer_obj.e_type == 'D':
                self.num_C_layer +=1
                self.layer_stack_simplified.append({'id':layer_obj.id,'electrical_type':layer_obj.e_type})
                
        self.layer_stack_simplified = sorted(self.layer_stack_simplified, key = lambda i: i['id'])       
        print(self.layer_stack_simplified)    
            
    def get_thick(self,layer_id):
        all_layer_info = self.layer_stack.all_layers_info
        layer = all_layer_info[layer_id]
        return layer.thick
    def get_z(self,layer_id):
        all_layer_info = self.layer_stack.all_layers_info
        layer = all_layer_info[layer_id]
        return layer.z
    
    
    
    def form_graph(self):
        self.doc_start_a_report()

        if self.ori_map == {}:
            self.find_ori = True
        else:
            self.find_ori = False
        isl_dict = {isl.name: isl for isl in self.layout_info}
        self.ele_lst = []
        self.hier_group_dict = {}
        self.comp_edge = []
        # add a step here to handle all layout island that are on the same layer then perform meshing all together.
        
        # STEP 1: Organize the layer_name and island_name
        
        for g in self.hier.isl_group:
            z = self.hier.z_dict[g.z_id]
            print('Z_level-', z,g.z_id)
            dz = self.get_thick(g.z_id)
            if not(g.z_id in self.layer_island_dict):
                self.layer_island_dict[g.z_id] = [g.name] # Add new list to collect island name
            else:
                self.layer_island_dict[g.z_id].append(g.name) # Add island name to layer
        
        # STEP 2: Process mesh elements for each layer and each island
        for layer_id in self.layer_island_dict:
            print(self.layer_island_dict)
            self.layer_mesh_table[layer_id] = LayerMesh()
            
            print("forming graph for layer:", layer_id)
            z = z = self.hier.z_dict[layer_id]
            layer_name = 'Layer_{}'.format(layer_id)
            
            for island_name in self.layer_island_dict[layer_id]:
                isl_mesh = TraceIslandMesh()
                isl = isl_dict[island_name]
                all_trace_copper = isl.elements
                all_net_on_trace = isl.child
                
            # add trace to the TraceIslandMesh object
                for trace_data in all_trace_copper:
                    x,y,width,height =  trace_data[1:5]
                    t_cell =RectCell(int(x),int(y),int(width),int(height)) 
                    isl_mesh.traces.append(t_cell)
                for net_data in all_net_on_trace:
                    name = net_data[5]
                    x,y,width,height =  net_data[1:5]
                    net_cell =RectCell(int(x),int(y),int(width),int(height)) 
                    isl_mesh.traces.append(net_cell)
                    if "L" in name: # lead type
                        isl_mesh.leads.append(net_cell)
                    elif "B" in name:
                        isl_mesh.pads.append(net_cell)
                    elif "D" in name:
                        isl_mesh.components.append(net_cell)
                    
                
                isl_mesh.form_hanan_mesh_table_on_island()
                isl_mesh.place_devices_and_components()
                self.layer_mesh_table[layer_id].add_table(island_name,isl_mesh)
            self.layer_mesh_table[layer_id].plot_all_mesh_island(name=layer_name)
            
            self.form_wire_frame(island=isl,isl_name = island_name,layer_name=layer_name,elv=z,dz=dz) # form wireframe for each island on each z level
            self.ele_lst.append(z)
            pos = {}
            # debug interface
            new_graph = deepcopy(self.graph)

            for n in self.graph.nodes: # Only plot the nodes on same level and save to pickle
                if self.graph.nodes[n]['island'] == layer_name:
                    pos[n] = self.graph.nodes[n]['locs'][0:2]
                else:
                    new_graph.remove_node(n)
            if self.debug:
                memfile = io.BytesIO()
                name = "digraph for layer -- " + str(layer_id)
                self.plot(mode=2,isl = g.name,pos = pos,graph = new_graph,save=False,mem_file = memfile)
                self.doc_handle_figure(memfile,name)
            self.ele_lst = list(set(self.ele_lst))
        for n in self.graph.nodes:
            self.pos[n] = self.graph.nodes[n]['locs'][0:2] # for 2D plotting purpose
        self.form_hierarchical_connections() # update bondwire new locations first
        for n in self.graph.nodes:
            if not n in self.pos3d:
                self.pos3d[n] = self.graph.nodes[n]['locs'] # for later graph rebuild
                self.nodes_dict_3d[tuple(self.graph.nodes[n]['locs'])] = n # for later graph rebuild
                
        input()

    def form_hierarchical_connections(self):
        '''
        Simply add the connection in the graph so the loop could be formed
        '''
        nets = self.comp_net_id
        for c in list(self.comp_dict.keys()):
            for e in c.net_graph.edges(data=True):
                e_type ='bondwire'
                if c.class_type =='comp':    
                    self.comp_edge.append([e[0],e[1],nets[e[0]],nets[e[1]]])
                    e_type = 'comp_edge'
                    node_type = self.graph.nodes[nets[e[1]]]['type']
                    if node_type == "trace":
                        self.contracted_nodes[nets[e[1]]] = nets[e[0]] # trace net to component net
                if c.class_type=='via':
                    self.comp_edge.append([e[0],e[1],nets[e[0]],nets[e[1]]])
                    e_type = 'via'
                if e_type == 'bondwire':
                    

                    n1 = nets[e[0]]
                    n2 = nets[e[1]]
                    p1 = self.graph.nodes[n1]['locs']
                    p2 = self.graph.nodes[n2]['locs']
                    left,right,top,bottom,thickness, z,ori = c.gen_ribbon(p1,p2)
                    rib_tc = TraceCell(left=left, right=right, bottom=bottom, top=top)
                    rib_tc.thick = thickness
                    rib_tc.z = z
                    rib_tc.dir = ori
                    rib_tc.struct = 'bw'
                    c = rib_tc.center()
                    # THIS Z UPDATE NEEDS TO BE UPDATED, FOR NOW USE ORIGINAL Z
                    z1=p1[2]
                    z2=p2[2]
                    #print ("Ribbon width height")
                    #print (ori)
                    #print (rib_tc.width_eval())
                    #print (rib_tc.height_eval())
                    #input()
                    #find connected trace to bondwire
                    if 'B' in e[0]: # find the net that connected to trace:
                        for ew in self.graph.edges(n1):
                            if ew[1]!=n2:
                                z1 = self.graph.nodes[ew[1]]['locs'][2]         
                    if 'B' in e[1]: # find the net that connected to trace:
                        for ew in self.graph.edges(n2):
                            if ew[1]!=n1:
                                z2 = self.graph.nodes[ew[1]]['locs'][2]        
                              
                    if ori == 1:
                        p1_new = (p1[0],c[1],z1)
                        p2_new = (p2[0],c[1],z2)
                    elif ori == 2:
                        p1_new = (c[0],p1[1],z1)
                        p2_new = (c[0],p2[1],z2)
                    self.nodes_dict_3d[p1_new]= n1
                    self.nodes_dict_3d[p2_new]= n2
                    self.pos3d[n1]=p1_new
                    self.pos3d[n2]=p2_new
                    self.pos[n1]=p1_new[0:2]
                    self.pos[n2]=p2_new[0:2]

                    rib_tc.bwn1=n1
                    rib_tc.bwn2=n2
                    
                    # BOND_WIRE
                    print ('bwribbon',rib_tc.eval_length()/1000, ori,'z',rib_tc.z)
                    self.ribbon_dict[(nets[e[0]], nets[e[1]])] = [rib_tc,thickness,z,ori] # get ribbon like data and store to a dictionary
                self.graph.add_edge(nets[e[0]], nets[e[1]],e_type = e_type,res_int=1)
    
    def form_bundles(self):
        
        self.net_graph = nx.Graph() # a graph for PEEC validation and loop evaluation

        
        for e in self.digraph.edges(data=True): # Loop through all edges in digraph (after contraction)
            # get edge data from graph
            #print(e)
            edata = self.graph.get_edge_data(e[0],e[1])
            #print ('digraph_data',edata)
            # for the new graph only
            if edata['e_type'] == 'comp_edge' or edata['e_type'] == 'via':
                data = {'type': edata['e_type'],'ori':None,'obj':None}
                #print('comp_edge',e[0],e[1])
                self.net_graph.add_edge(e[0],e[1],data=data,res= 1e-12,ind=1e-12,len=10) # add original compedge to loop_graph to ensure closed loop later         
            if edata == None : # self_loop from merging edge, cant be found
                bw_net = self.contracted_nodes[e[1]] # get the contracted node
                edata = self.graph.get_edge_data(bw_net,e[0]) # bw_net is already contracted *
                #print('merged_edge',e[0],e[1])

                data = {'type': edata,'ori':None,'obj':None}
                self.net_graph.add_edge(n1,n2,data=data,res= 1e-12,ind=1e-12,len=10) # add original compedge to loop_graph to ensure closed loop later
                if edata == None or edata == 'comp_edge':
                    continue
         
            locx = [self.pos[e[0]][0],self.pos[e[1]][0]]
            locy = [self.pos[e[0]][1],self.pos[e[1]][1]]
            
            # handle traces
        
            if edata['e_type'] == 'trace':
                ptrace = edata['p_trace'] # Get parent trace
                # based on edge data to check direction
                # NOTE TO SELF: tbl_isl_node_connection needs to map isl --> net --> directions_dict (SIMPLIFY LOOPS)
                for dir in self.tbl_isl_node_connection[e[0]]: # from the first edge find the direction of the second edge
                    if self.tbl_isl_node_connection[e[0]][dir] == e[1]:
                        # Need to change trace cell bottom and top positions accordingly
                        if dir == 'N':
                            trace = deepcopy(ptrace)
                            trace.dir = 2
                            trace.bottom = min(locy)
                            trace.top = max(locy)
                            trace.width_eval()
                            trace.height_eval()
                            self.y_cells.append(trace)
                        elif dir == 'S':
                            trace = deepcopy(ptrace)
                            trace.dir = -2
                            trace.bottom = min(locy)
                            trace.top = max(locy)
                            trace.width_eval()
                            trace.height_eval()
                            self.y_cells.append(trace)
                        
                        # Need to update trace cell left and right positions accordingly 
                        elif dir == 'E':
                            trace = deepcopy(ptrace)
                            trace.dir = 1
                            trace.left = min(locx)
                            trace.right = max(locx)
                            trace.width_eval()
                            trace.height_eval()
                            self.x_cells.append(trace)
                        elif dir == 'W':
                            trace = deepcopy(ptrace)
                            trace.dir = -1
                            trace.left = min(locx)
                            trace.right = max(locx)
                            trace.width_eval()
                            trace.height_eval()
                            self.x_cells.append(trace)
                        self.tc_to_edges_init[trace] = e # map the tracecell to edge for R, L update later  
            elif edata['e_type'] == 'bondwire':
                #bw_net = self.contracted_nodes[e[0]] 
                n1 = e[0]
                n2 = e[1]
                try:
                    rib_tc,thick,z,ori = self.ribbon_dict[(n1,n2)]
                except:
                    rib_tc,thick,z,ori = self.ribbon_dict[(n2,n1)]

                rib_tc.struct = 'bw'
                p1 = self.graph.nodes[n1]['locs']
                p2 = self.graph.nodes[n2]['locs']
                dir = ori
                if ori ==1:
                    if p1[0] > p2[0]:
                        dir = -1
                    #self.x_cells.append(rib_tc)
                    rib_tc.dir =1
                    if not (p1[0],p2[0]) in self.x_wires:
                        self.x_wires[p1[0],p2[0]] = [rib_tc]
                    else:
                        self.x_wires[p1[0],p2[0]].append(rib_tc)

                    
                elif ori == 2:
                    if p1[1] > p2[1]:
                        dir = -2 
                    rib_tc.dir =2       
                    #self.y_cells.append(rib_tc)
                    if not (p1[1],p2[1]) in self.y_wires:
                        self.y_wires[p1[1],p2[1]] = [rib_tc]
                    else:
                        self.y_wires[p1[1],p2[1]].append(rib_tc)
                rib_tc.dir = dir
                ori_str = ('h' if ori == 1 else 'v')  
                data = {'type': edata,'ori':ori_str,'obj':rib_tc}
                #self.net_graph.add_edge(n1,n2,data=data,res= 1e-12,ind=1e-12) # add original compedge to loop_graph to ensure closed loop later
                #print ('bw_edge',n1,n2)
                self.tc_to_edges_init[rib_tc] = e # map the tracecell to edge for R, L update later  

                #print("get ribbon data")
        '''
        print("number of elements", len(self.x_cells))
        
        for tx in self.x_cells:
            print (tx,tx.dir)
        for ty in self.y_cells:
            print (ty,ty.dir)
        '''
        self.x_bundles = {} # Store all bundle in x to compute in parallel
        self.y_bundles = {} # Store all bundle in y to compute in parallel
        # Split on Trace Level only
        self.split_bundles(direction = 'x')
        self.split_bundles(direction = 'y')

    def split_bundles(self,direction = 'x'):
        '''
        Add a chunk of code here to define ground return path or just extract the current of the ground mesh and invert the matrix
        # by default set -1 direction to G
        '''
        # step 1:
        # Loop through all horizontal trace cells created from the previous step
        # Collect x/y locs to split bundle.
        y_locs = []
        x_locs = []
        x_locs_edge = {}
        y_locs_edge = {}
        if direction == 'x': # 1N operations
            for tx in self.x_cells:
                c = tx.center()
                y_locs.append(c[1])
                x_locs += [tx.left,tx.right]
                e = self.tc_to_edges_init[tx] # get the initial trace cell
                x_locs_edge[(tx.left,tx.right,c[1])] = e # map the left, right, y to edge  
            # Step 2: Now search through each trace and split them
            y_locs = list(set(y_locs))
            y_locs.sort()
            x_locs = list(set(x_locs))
            x_locs.sort()
            xcells = deepcopy(self.x_cells)
            Nx = len(xcells)
            self.x_cells = [] # empty the original xcell to rebuild
            flags = [0 for x in range(Nx)] # make a list of flag to know which cell is splitted 
            rm_bundles = {} # To flag 1 to the short bundles that can be removed
            # Form x locs bundles
            for i in range(len(x_locs)-1): 
                dx = abs(x_locs[i]-x_locs[i+1])
                #self.x_bundles[(x_locs[i],x_locs[i+1])]=[]
                
                if dx > 100 : # 0.5 mm
                    self.x_bundles[(x_locs[i],x_locs[i+1])]=[]
                    rm_bundles[(x_locs[i],x_locs[i+1])] = 0
                else:
                    rm_bundles[(x_locs[i],x_locs[i+1])] = 1
                
            for y_loc in y_locs:
                
                for i in range(Nx):
                    tx = xcells[i]
                    c = tx.center()
                    y = c[1]
                    xmin = tx.left
                    xmax = tx.right
                    z= tx.z
                    if  y!=y_loc :
                        continue
                    if flags[i] != 1:
                        flags[i] = 1
                        for j in range(len(x_locs)-1):
                            if rm_bundles[x_locs[j],x_locs[j+1]] == 1:
                                if x_locs[j]>= xmin and x_locs[j]< xmax:
                                    self.contracted_edges.append([(x_locs[j],y,z),(x_locs[j+1],y,z)])
                                continue
                            if x_locs[j]>= xmin and x_locs[j]< xmax: # Overlapped regions
                                new_tx = deepcopy(tx)
                                new_tx.left = x_locs[j]
                                new_tx.right = x_locs[j+1]
                                new_tx.width_eval()
                                new_tx.height_eval()
                                self.x_bundles[(new_tx.left,new_tx.right)].append(new_tx)
                                #self.tc_to_edges_splitted[new_tx] = x_locs_edge[(x_locs[j],x_locs[j+1],y)]
                                self.x_cells.append(new_tx)
                            
            
        else:

            rm_bundles = {} # To flag 1 to the short bundles that can be removed
            for ty in self.y_cells:
                c = ty.center()
                x_locs.append(c[0])
                y_locs += [ty.bottom,ty.top]
                e = self.tc_to_edges_init[ty] # get the initial trace cell
                y_locs_edge[(ty.bottom,ty.top,c[0])] = e # map the top, bot, x to edge 
            # Step 2: Now search through each trace and split them
            y_locs = list(set(y_locs))
            y_locs.sort()
            x_locs = list(set(x_locs))
            x_locs.sort()
            ycells = deepcopy(self.y_cells)
            Ny = len(ycells)
            self.y_cells = [] # empty the original xcell to rebuild
            flags = [0 for y in range(Ny)] # make a list of flag to know which cell is splitted 
            # Form x locs bundles
            for i in range(len(y_locs)-1): # Only form a bundle if the length is greater than 0.1 mm
                dy = abs(y_locs[i]-y_locs[i+1])
                #self.y_bundles[(y_locs[i],y_locs[i+1])]=[] 
                if dy > 100:
                    self.y_bundles[(y_locs[i],y_locs[i+1])]=[] 
                    rm_bundles[(y_locs[i],y_locs[i+1])] = 0
                else:
                    rm_bundles[(y_locs[i],y_locs[i+1])] = 1
                

            
            for x_loc in x_locs:
                for i in range(Ny):
                    ty = ycells[i]

                    c = ty.center()
                    x = c[0]
                    ymin = ty.bottom
                    ymax = ty.top
                    z=ty.z
                    if  x!=x_loc :
                        continue
                    if flags[i] != 1:
                        flags[i] = 1
                        for j in range(len(y_locs)-1):
                            if rm_bundles[y_locs[j],y_locs[j+1]] == 1:
                                if y_locs[j]>= ymin and y_locs[j]< ymax:
                                    self.contracted_edges.append([(x,y_locs[j],z),(x,y_locs[j+1],z)])

                                continue
                            if y_locs[j]>= ymin and y_locs[j]< ymax: # Overlapped regions

                                new_ty = deepcopy(ty)
                                new_ty.bottom = y_locs[j]
                                new_ty.top = y_locs[j+1]
                                new_ty.width_eval()
                                new_ty.height_eval()
                                if ty.struct == 'bw': # in case a bondwire is splitted
                                    # get x,y of old node:
                                    x1, y1 = self.pos[ty.bwn1]
                                    x2, y2 = self.pos[ty.bwn2]
                                    if y1 > y2:
                                        c_top = y1
                                        c_bot = y2
                                        if c_top != new_ty.top:
                                            new_ty.bwn1 = None # reset for later
                                        elif c_bot != new_ty.bottom:
                                            new_ty.bwn2 = None
                                    else:
                                        c_top = y2
                                        c_bot = y1
                                        if c_top != new_ty.top:
                                            new_ty.bwn2 = None  # reset for later
                                        elif c_bot != new_ty.bottom:
                                            new_ty.bwn1 = None
                                self.y_bundles[(new_ty.bottom,new_ty.top)].append(new_ty)
                                #self.tc_to_edges_splitted[new_ty] = y_locs_edge[(y_locs[j],y_locs[j+1],x)]
                                self.y_cells.append(new_ty)

                # Get the center point go get y locs
            #self.plot_xy_trace(mode=0,dir = 'X')
            #self.plot_xy_trace(mode=0,dir = 'Y')


    def plot_xy_trace(self, mode = 0, dir = 'X'):
        # if mode = 0 then plot here, if = 1 then save to pickles
        
    
        if dir == "X":
            fig1 = plt.figure("3d xbundles")
            ax1 = plt.axes(projection='3d')
            
            plt.xlim(-10000,50000)
            plt.ylim(-10000,50000)

            for tx in self.x_cells:
    
                xs = [tx.left, tx.right]
                ys = [tx.bottom,tx.top]
                X,Y = np.meshgrid(xs,ys)
                Z = np.ones((2,2))
                Z *= tx.z
                
                if tx.dir == 1:
                    ax1.plot_surface(X,Y,Z, color = 'red')  
                elif tx.dir == -1:
                    ax1.plot_surface(X,Y,Z, color = 'blue')            
            if mode == 0:
                plt.show()
                #plt.savefig("fig/horizontal bundles")
            else:
                m_name = './pickles/x_bundle'  + '.p'
                #self.save_plot_pickle(ax1,m_name)    
        elif dir == "Y":
            fig1 = plt.figure("3d ybundles")
            ax1 = plt.axes(projection='3d')
            plt.xlim(-10000,50000)
            plt.ylim(-10000,50000)

            for tx in self.y_cells:
                xs = [tx.left, tx.right]
                ys = [tx.bottom,tx.top]
                X,Y = np.meshgrid(xs,ys)
                Z = np.ones((2,2))
                Z *= tx.z
                if tx.dir == 2:
                    ax1.plot_surface(X,Y,Z, color = 'red')
                elif tx.dir == -2:
                    ax1.plot_surface(X,Y,Z, color = 'blue')
            if mode == 0:
                plt.show()

                #plt.savefig("fig/vertical bundles")
            else:
                m_name = './pickles/y_bundle'  + '.p'
                #self.save_plot_pickle(ax1,m_name)
        elif dir == "3D":
            fig3 = plt.figure("3d bundles")
            ax3d = plt.axes(projection='3d')
            xxs = []
            xys = []
            xzs = []
            yxs = []
            yys = []
            yzs = [] 
            for tx in self.x_cells:    
                tx.width_eval()
                tx.height_eval()
                c = tx.center()  
                ax3d.plot3D(c[0])

    def find_ground_wire(self, ori = 0, traces=[]): # NEED TO RETHINK THIS 
        # find the ground wire of the group of each bundle
        # ori = 0 or 1 for horizontal vs vertical cases
        # traces in a bundle
        dir_to_trace = {}
        wire_type = {}
        sig_w = []
        gr_w = []
        # first search through all traces and store them in bucket, where each bucket is a direction
        for tr in traces:
            if not tr.dir in dir_to_trace:
                dir_to_trace[tr.dir]=[]
            else:
                dir_to_trace[tr.dir].append(tr)
        dir_list = list(dir_to_trace.keys())
        if len(dir_list)==1: # means that they only have one dir
            for tr in traces:
                wire_type[tr] = 'S'
        else:
            # by default we set the wires in negative direction to be ground
            # if it is bondwire, we should set them to be signal type

            for tx in traces:

                if tx.dir <0 :
                    wire_type[tx] = 'G' # 'S" set all to signal
                    gr_w.append(tx)
                    #sig_w.append(tx)
                    
                else:
                    wire_type[tx] = 'S'
                    sig_w.append(tx)

            if len(sig_w) == 0:
                sig_w = gr_w
                gr_w = [] # Test no ground wire
                for w in wire_type:
                    wire_type[w] = 'S' # switch from G to S
        return wire_type
    def add_node_to_3d(self,pos):

        if pos in self.nodes_dict_3d:
            node = self.nodes_dict_3d[pos]
        else:
            node = self.new_nodes_id
            self.nodes_dict_3d[pos] = node
            self.pos3d[node] = pos
            self.new_nodes_id += 1
            self.net_graph.add_node(node, locs=pos)
        if node == 28:
            print (node , pos)
        return node
    def rebuild_graph_add_edge(self,tc,n1,n2,R,L,etype='fw'):
        #print(n1,n2,"R",R,"L",L)
        #print(self.comp_net_id)
        z = tc.z

        if tc.struct=='trace': # for normal bundles
            #print("TRACE:", tc)
            if abs(tc.dir) == 1:
                pos1 = (tc.left,tc.center()[1],z)
                pos2 = (tc.right,tc.center()[1],z)
                n1 = self.add_node_to_3d(pos1)
                n2 = self.add_node_to_3d(pos2)
                data= {'type':etype,'ori':'h','obj':tc}
                self.net_graph.add_edge(n1,n2,data=data,res=R,ind=L,len=10)
            if abs(tc.dir) == 2:
                pos1 = (tc.center()[0], tc.bottom,z)
                pos2 = (tc.center()[0],tc.top,z)
                n1 = self.add_node_to_3d(pos1)
                n2 = self.add_node_to_3d(pos2)
                data= {'type':etype,'ori':'v','obj':tc}

                self.net_graph.add_edge(n1,n2,data=data,res=R,ind=L,len=10)
        elif tc.struct == 'bw':
            if tc.bwn1!= None and tc.bwn2 != None:
                n1 = tc.bwn1
                n2 = tc.bwn2
            elif tc.bwn1 == None and tc.bwn2!= None: # n2 is on the graph n1 is not
                n2 = tc.bwn2
                pos2 = self.pos3d[n2] # n2 is on the old graph
                if abs(tc.dir) == 1: # horizontal case
                    find_right = False
                    if pos2[0] == tc.left:
                        find_right = True
                    if find_right:
                        pos1 = (tc.right, tc.center()[1], z)
                    else: # find left
                        pos1 = (tc.left, tc.center()[1], z)
                    n1 = self.add_node_to_3d(pos1)
                if abs(tc.dir) == 2:
                    find_top = False
                    if pos2[1] == tc.bottom:
                        find_top = True
                    if find_top:
                        pos1 = (tc.center()[0], tc.top, z)
                    else: # find bottom
                        pos1 = (tc.center()[0], tc.bottom, z)
                n1 = self.add_node_to_3d(pos1)
            elif tc.bwn1 != None and tc.bwn2 == None:  # n1 is on the graph n2 is not
                n1 = tc.bwn1
                pos1 = self.pos3d[n1]  # n2 is on the old graph
                if abs(tc.dir) == 1:  # horizontal case
                    find_right = False
                    if pos1[0] == tc.left:
                        find_right = True
                    if find_right:
                        pos2 = (tc.right, tc.center()[1], z)
                    else:  # find left
                        pos2 = (tc.left, tc.center()[1], z)
                    n2 = self.add_node_to_3d(pos2)
                if abs(tc.dir) == 2:
                    find_top = False
                    if pos1[1] == tc.bottom:
                        find_top = True
                    if find_top:
                        pos2 = (tc.center()[0], tc.top, z)
                    else:  # find bottom
                        pos2 = (tc.center()[0], tc.bottom, z)
                n2 = self.add_node_to_3d(pos2)

            ori = ('h' if tc.dir==1 else 'v')
            data= {'type':etype+'_bw','ori':ori,'obj':tc}
            #print (data,R,L)
            # THIS IS TO EXCLUDE BW LOOP (TESTING) 
            if self.short_bw:
                R = 1e-12
                L = 1e-12
            self.net_graph.add_edge(n1,n2,data=data,res=R,ind=L,len=10)
        return n1,n2
    def rebuild_graph(self,loop_obj,mode = ""):
        m_dict = {}
        l_eval_dict={}
        rem_dict = {} # to make sure there is no overlapping in mutual pair
        rebuild_graph_info = {} # for debugging purpose and back-annotation of net-graph figure.
        
        if mode == "eval_ground_imp":
            print(loop_obj.L_loop)
            for tc in loop_obj.tc_to_id_gr:
                id1 = loop_obj.tc_to_id_gr[tc]
                R1 = abs(loop_obj.R_loop[id1,id1])
                L1 = abs(loop_obj.L_loop[id1,id1])
                # ADD self RL values
                
                n1,n2 = self.rebuild_graph_add_edge(tc,0,0,R1,L1,etype = 'fw')
                rebuild_graph_info[(n1,n2)]="type:{} R: {} L:{} ".format('fw',R1,L1)

                m_dict[tc] = (n1,n2)
                l_eval_dict[tc] = L1
            
            # handle mutual and k coefficients
            check_list={}
            
            for tc1 in loop_obj.tc_to_id_gr:
                id1 = loop_obj.tc_to_id_gr[tc1]
                if id1 == -1:
                    continue    
                e1 = m_dict[tc1]
                L1 = l_eval_dict[tc1]
                for tc2 in loop_obj.tc_to_id_gr:
                    id2 = loop_obj.tc_to_id_gr[tc2]
                    if id1 == id2 or id2 == -1:
                        continue
                    else:
                        e2 = m_dict[tc2]
                        L2= l_eval_dict[tc2]
                        if not (id1,id2) in check_list:
                            check_list[(id1,id2)]= 1
                            check_list[(id2,id1)]= 1
                            print(id1,id2)
                            
                            self.mutual_pair[(e1,e2)] = (0, (loop_obj.L_loop[id1,id2]))
                            k = loop_obj.L_loop[id1,id2]/math.sqrt(L1*L2)
                            self.mutual_coup_coeff_pair[(e1,e2)] = k
                            m_dict[e1] = e2
                            m_dict[e2] = e1
            #input()
        elif mode == "eval_normal":

            for tc in loop_obj.tc_to_id:
                id1 = loop_obj.tc_to_id[tc]
                if id1!= -1:
                    # find n1 n2 based on old locs
                    R1 = abs(loop_obj.R_loop[id1,id1])
                    L1 = abs(loop_obj.L_loop[id1,id1])
                    # ADD self RL values
                    
                    n1,n2 = self.rebuild_graph_add_edge(tc,0,0,R1,L1,etype = 'fw')
                    rebuild_graph_info[(n1,n2)]="type:{} R: {} L:{} ".format('fw',R1,L1)

                    m_dict[tc] = (n1,n2)
                else:
                    new_n2 = False

                    if abs(tc.dir) == 1:
                        pos1 = (tc.left,tc.center()[1],tc.z)
                        pos2 = (tc.right,tc.center()[1],tc.z)
                        
                    elif abs(tc.dir) ==2 :
                        pos1 = (tc.center()[0], tc.bottom,tc.z)
                        pos2 = (tc.center()[0],tc.top,tc.z)
                    
                    if pos1 in self.nodes_dict_3d:
                        n1 = self.nodes_dict_3d[pos1]
                    else:
                        n1 = self.new_nodes_id
                        new_n2 = True
                    if pos2 in self.nodes_dict_3d:
                        n2 = self.nodes_dict_3d[pos2]
                    else:
                        if new_n2:
                            n2 = self.new_nodes_id+1
                        else:
                            n2 = self.new_nodes_id
                    n1,n2 = self.rebuild_graph_add_edge(tc,n1,n2,1e-12,1e-12,etype='return')
                    rebuild_graph_info[(n1,n2)]="type:{} R: {} L:{} ".format('return',1e-12,1e-12)

            if loop_obj.num_loops > 1: # ADD mutual value
                for tc1 in loop_obj.tc_to_id:
                    id1 = loop_obj.tc_to_id[tc1]
                    if id1 == -1:
                        continue    
                    e1 = m_dict[tc1]
                    for tc2 in loop_obj.tc_to_id:
                        id2 = loop_obj.tc_to_id[tc2]
                        if id1 == id2 or id2 == -1:
                            continue
                        else:
                            e2 = m_dict[tc2]

                            if not (e1 in m_dict):
                                #if not self.short_bw:
                                self.mutual_pair[(e1,e2)] = (0, (loop_obj.L_loop[id1,id2]))
                                #self.mutual_pair[(e2,e1)] = (loop_obj.R_loop[id1,id2], loop_obj.L_loop[id1,id2]) # ensure mutual value is considered
                                m_dict[e1] = e2
                                m_dict[e2] = e1

                        #self.rebuild_graph_add_edge(tc1,n3,n4,R2,L2) # add the self value for this trace cell
        if self.debug:
            self.doc_report.add_paragraph("Rebuild graph info:")
            self.doc_handle_rebuild_graph_info(rebuild_graph_info)
    
    def handle_contracted_edge(self):
        con_edge = []
        for e in self.contracted_edges: # store location of each node
            if e[0] in self.nodes_dict_3d:
                n1 = self.nodes_dict_3d[e[0]]
            else:
                n1 = self.new_nodes_id
                self.nodes_dict_3d[e[0]] = n1
                self.pos3d[n1] = e[0] 
                self.net_graph.add_node(n1,locs = e[0])
                self.new_nodes_id+=1

            if e[1] in self.nodes_dict_3d:
                n2 = self.nodes_dict_3d[e[1]]
            else:
                n2 = self.new_nodes_id
                self.nodes_dict_3d[e[1]] = n2
                self.pos3d[n2] = e[1] 
                self.net_graph.add_node(n2,locs = e[1])
                self.new_nodes_id+=1
            data = {'type': 'small_bundle','ori':None,'obj':None}
            self.net_graph.add_edge(n1,n2,data=data,res= 1e-12,ind=1e-12,len=10)
        self.net_2d_pos = {} # to store 2d location for 2d net graphing purpose
        for n in self.net_graph.nodes:
            try:
                self.net_2d_pos[n] = self.net_graph.nodes[n]['locs'][0:2] # for 2D plotting purpose
            except: # must be in the 3d pos, added previously
                #print ("no 3d pos",n)
                self.net_2d_pos[n] = self.pos3d[n][0:2] # for 2D plotting purpose
        n_list = list(self.net_2d_pos.keys())
        n_list.sort()
        #for n in n_list:
        #    print(n,self.pos3d[n])
        

    def build_PEEC_graph(self):
        self.new_nodes_id = max(self.graph.nodes) +1 # start from the max value of sefl.graph.nodes
        self.PEEC_graph = nx.Graph()
        start_node_id = 1 # index from 1
        loop_eval_model = LoopEval("PEEC") # Use this to compute all filaments if PEEC technique is used
        edge_to_trace_dict = {}
        loc_to_node = {}
        edge_to_m_id = {}
        short_bw = False
        for e in self.net_graph.edges(data = True):

            edata = e[2]['data']
            type=e[2]['data']['type']
            trace_obj = edata['obj']
            
            if trace_obj!= None:
                trace_mesh = loop_eval_model.add_trace_cell(trace_obj)
                if short_bw:
                    if type =='fw_bw':
                        self.PEEC_graph.add_edge(e[0],e[1],data=None)
                        continue
                # Include the trace mesh object here to link the edges with the filaments
                self.PEEC_graph.add_edge(e[0],e[1],data = trace_mesh, Zdict={})
            else:
                self.PEEC_graph.add_edge(e[0],e[1],data=None)
        # This will update the R, L and mutual inductance between all filaments
        loop_eval_model.form_partial_impedance_matrix()
        loop_eval_model.update_mutual_mat()
        # Quick Access to loop obj R and L Matrices
        R_Mat = loop_eval_model.R_Mat
        L_Mat = loop_eval_model.L_Mat

        # Update PEEC graph
        for e in self.PEEC_graph.edges(data= True):
            trace_mesh = e[2]['data']
            Z_dict = {}
            if trace_mesh == None:
                continue
            for el in trace_mesh.elements:
                Zkey = "B{0}".format(el.id)
                Rval = R_Mat[el.id,el.id]
                Lval = L_Mat[el.id,el.id]
                Z_dict[Zkey]=Rval + Lval*1j
            e[2]['Zdict']=Z_dict
        self.M_PEEC = {}
        for i in range(loop_eval_model.tot_els):
            for j in range(loop_eval_model.tot_els):
                if i == j or (j,i) in self.M_PEEC:
                    continue
                if L_Mat[i,j]<0:
                    input("num error")
                self.M_PEEC[(i,j)] = L_Mat[i,j]

    def solve_all_bundles(self):
        # UPDATE WITH MULTI PROCESSING ***
        # FOR ECCE simply rebuild the graph, but once we need to reverse this algorithm to make it more efficient 
        # SOLVE X BUNDLES
        
        self.new_nodes_id = max(self.graph.nodes) +1 # start from the max value of sefl.graph.nodes
        bundle_id = 0
        x_loops = []

        for bundle in self.x_bundles:
            print("Updating BUNDLE: ", bundle)
            loop_model = LoopEval('x_'+str(bundle_id))
            loop_model.frequency = self.frequency*1000
            #loop_model.mesh_method='nonunifom'
            bundle_id+=1
            wire_type = self.find_ground_wire(ori = 0, traces = self.x_bundles[bundle])
            for trace in self.x_bundles[bundle]:
                loop_model.add_trace_cell(trace,el_type = wire_type[trace])
            

            x_loops.append(loop_model)

        # SOLVE Y BUNDLES
        '''
        loop_model.solve_linear_systems()
        self.rebuild_graph(loop_model)
        '''
        y_loops = []
        bundle_id = 0
        for bundle in self.y_bundles:
            #print("SOLVING BUNDLE: ", bundle)
            loop_model = LoopEval('y_'+str(bundle_id))
            loop_model.frequency = self.frequency*1000
            #loop_model.mesh_method='nonunifom'
            bundle_id+=1
            wire_type = self.find_ground_wire(ori = 1, traces = self.y_bundles[bundle])
            for trace in self.y_bundles[bundle]:
                loop_model.add_trace_cell(trace,el_type = wire_type[trace])
            

            y_loops.append(loop_model)
        bundle_id = 0
        
        for bundle in self.x_wires:
            loop_model = LoopEval('x_wire_'+str(bundle_id))

            bundle_id+=1

            wire_type = self.find_ground_wire(ori = 0, traces = self.x_wires[bundle])
            for trace in self.x_wires[bundle]:
                loop_model.add_trace_cell(trace,el_type = wire_type[trace])
            
            x_loops.append(loop_model)
        bundle_id = 0
        
        for bundle in self.y_wires:
            loop_model = LoopEval('y_wire_'+str(bundle_id))

            bundle_id+=1

            wire_type = self.find_ground_wire(ori = 1, traces = self.y_wires[bundle])
            for trace in self.y_wires[bundle]:
                loop_model.add_trace_cell(trace,el_type = wire_type[trace])

            y_loops.append(loop_model)
        

        all_loops = x_loops+y_loops
        #update_all_mutual_ele(all_loops)
        self.all_loops= all_loops

        # SOLVE EACH BUNDLE SEPARATEDLY, POSSIBLE FOR PARRALLEL RUN
        if self.debug:
            self.doc_report.add_heading("Checking for nets and nodes:",1)
            self.doc_report.add_heading("node dictionary",2)
            for n in self.nodes_dict_3d:
                self.doc_report.add_paragraph("Node 3D location: {} --- Node ID: {}".format(n,self.nodes_dict_3d[n]))
            self.doc_report.add_heading("component to node dictionary",2)
            for c in self.comp_net_id:
                self.doc_report.add_paragraph("Component name: {} --- Component Node id: {}".format(c,self.comp_net_id[c]))
            self.doc_report.add_heading("Checking bundle creation and evaluation result",1)
            self.doc_report.add_paragraph("Total number of horizontal bundles: {} ".format(len(x_loops)))
            self.doc_report.add_paragraph("Total number of vertical bundles: {}".format(len(y_loops)))
        
        #self.all_loops=solve_loop_models_parallel(self.all_loops)
        mode ='equation'
        if mode == 'equation':
            meshing_algorithm = 'nonuniform'
        elif mode == 'regression':
            meshing_algorithm = 'uniform_fixed_width'
        eval_mode = 'eval_ground_imp'
        if eval_mode =='eval_ground_imp':
            decoupled = True
        elif eval_mode == 'eval_normal':
            decoupled = False
        for loop_model in self.all_loops:
            print("evaluating ... ", loop_model.name)
            loop_model.mode = mode
            
            loop_model.form_mesh_traces(mesh_method=meshing_algorithm)
            loop_model.form_partial_impedance_matrix()
            loop_model.update_mutual_mat()
            loop_model.form_mesh_matrix()
            loop_model.update_P()
            loop_model.solve_linear_systems(decoupled=decoupled)
            #if 'wire' in loop_model.name:
            #    print(loop_model.L_loop)
            if self.debug:
                self.doc_export_report_for_each_loop(loop_model)
            #    P_df = pd.DataFrame(data=loop_model.P)
            #    P_df.to_csv("P_mat_{}.csv".format(loop_model.name))
            self.rebuild_graph(loop_model,mode=eval_mode)
            


        #print (self.mutual_pair)
        self.handle_contracted_edge()
        edge_list = []
        selected_node = []
        for e in self.net_graph.edges(data = True):
            edata = e[2]['data']
            
            if 'fw' in edata['type']:
                edge_list.append((e[0],e[1]))
                selected_node.append(e[0])
                selected_node.append(e[1])
                

            elif 'return' in edata['type']:
                edge_list.append((e[0],e[1]))
                selected_node.append(e[0])
                selected_node.append(e[1])
        
        
        
        # TEST new graph
        #print ("check path")
        #paths = nx.all_simple_paths(self.net_graph,5,14)
        #for p in paths:
        #    print(p)
        #input()
        if self.debug:
            memfile = io.BytesIO()
            self.plot(mode=5,save=True,mem_file=memfile)
            self.doc_handle_figure(memfile=memfile,fig_heading = "Netlist Graph")
            self.doc_print_net_list_line_by_line()
            self.doc_save_a_report('./debug_report.docx')  
    '''The below functions are used to format the report for debugging purpose'''
    
    
    def doc_start_a_report(self):
        # This is a debug report for this model. It would include init-layout, mesh structure
        if self.debug:    
            self.doc_report= Document()
            self.doc_report.add_heading("Automated Report for Loop-Based Extraction",0)
            today = date.today()
            d_format = today.strftime("%m/%d/%y")
            first_line= "This is an automated report for the current layout in debug mode. Report generated on {}".format(d_format)
            self.doc_report.add_paragraph(first_line)  
    
    
    def doc_print_net_list_line_by_line(self):
        for e in self.net_graph.edges(data = True):
            edata = e[2]['data']
            eval = e[2]
            if 'fw' in edata['type']:
                line = "{} -- {} -- R: {} Ohm, L: {} Henry \n".format(e[0],e[1],eval['res'],eval['ind'])
                self.doc_report.add_paragraph(line)    
        # compute K
        for m in self.mutual_pair:
            e1 = m[0]
            e2 = m[1]
            line = "edge 1 {} -- edge 2 {}, K {} ,M {}".format(e1,e2,self.mutual_coup_coeff_pair[m],self.mutual_pair[m])
            self.doc_report.add_paragraph(line)    

            
    
    def doc_export_report_for_each_loop(self,loop):
        self.doc_report.add_heading("Analysis for loop: {} \n".format(loop.name),2)
        text= loop.export_loop(mode = 1)
        self.doc_report.add_paragraph(text)
        self.doc_write_matrix_to_report(np.abs(loop.R_loop),"R_Loop of "+ loop.name)
        self.doc_write_matrix_to_report(np.abs(loop.L_loop),"L_loop of "+ loop.name)
        #self.doc_write_matrix_to_report(loop.I,"I_Matrix of "+loop.name)

    def doc_save_a_report(self,file):
        self.doc_report.save(file)

    def doc_write_matrix_to_report(self,matrix,name):
        nx,ny = list(matrix.shape)
        self.doc_report.add_heading("Matrix " + name , 3)
        text = "Matrix {} with {} rows and {} columns".format(name,nx,ny)
        self.doc_report.add_paragraph(text)
        
        table = self.doc_report.add_table(rows=1,cols = ny+1)
        table.style ='Table Grid'

        row = table.rows[0].cells
        for iy in range(ny+1): # first row
            if iy == 0:
                continue
            id = iy -1
            row[iy].text = str(id)
        for ix in range(nx):
            row = table.add_row().cells
            row[0].text = str(ix)
            for iy in range(ny):
                row[iy+1].text = str(matrix[ix,iy])
    def doc_handle_rebuild_graph_info(self,loop_rebuild):
        for k in loop_rebuild:
            text = "{}---{}".format(k[0],k[1]) + " " + loop_rebuild[k]
            self.doc_report.add_paragraph(text)
    def doc_handle_figure(self,memfile,fig_heading='fig'):
        # store the iostream figure to doc
        self.doc_report.add_heading(fig_heading,2)
        self.doc_report.add_picture(memfile,width = Inches(4),height=Inches(4))
    '''The upper functions are used to format the report for debugging purpose'''

    def graph_to_circuit_transfomation(self):
        for e in self.contracted_edges:
            try:
                self.digraph = nx.contracted_nodes(self.digraph,e[0],e[1],self_loops = False)
            except:
                self.digraph = nx.contracted_nodes(self.digraph,e[1],e[0],self_loops = False)
        #plt.figure("digraph_contracted")
        #nx.draw(self.digraph,self.pos,with_labels=True)
        #plt.savefig('digraph_contracted.png')
    
    def add_and_update_nodes(self,locs,isl_name,tc,type='trace'):
        if not locs in self.tbl_isl_loc_net[isl_name]:
            self.graph.add_node(self.node_id, locs = locs ,island = isl_name, type = type,parent = [tc])
            self.tbl_isl_loc_net[isl_name][locs] = self.node_id
            self.node_id +=1
    # Support function to reduce complexity
    def add_node_tracecell(self,tc,x,y,isl_name):
        '''
        add a hierachical loc on H or V trace_cell 
        --------------------------
        -        *               -
        --------------------------
        '''
        z= tc.z
        if self.ori_map[tc.name] == 'H':
            locx =(x,tc.bottom+tc.height/2,z)
            self.graph.add_node(self.node_id,locs=locx,island = isl_name, type = 'trace',parent = [tc])
            self.tbl_isl_loc_net[isl_name][locx]= self.node_id
            self.node_id+=1 
        elif self.ori_map[tc.name] == 'V':
            locy =(tc.left+tc.width/2,y,z)
            self.graph.add_node(self.node_id,locs=locy,island = isl_name, type = 'trace',parent = [tc])
            self.tbl_isl_loc_net[isl_name][locy]= self.node_id

            self.node_id+=1
        ''' HANDLE PLANAR CASE LATER
        else: # planar case
            self.graph.add_node(self.node_id,locs=[x,y,z],island = isl_name, type = 'trace',parent = [tc])
            isl_nodes.append(self.node_id)
            self.node_id+=1
        '''
    def add_node_float(self, loc, net,isl_name):
        '''
        add a floating net from device terminal that is not on the trace mesh (e.g for a device signal pin or via)
        '''
        self.graph.add_node(self.node_id,locs=loc,island = isl_name, type = 'hier',parent = None) # double check this in 3D
        self.tbl_isl_loc_net[isl_name][loc] = self.node_id
        self.comp_net_id[net] = self.node_id
        self.node_id+=1
    def create_strace(self, edge_data,isl_name,elv,dz):
        '''
        create a simple trace and add nodes  ("H" or "V")
        '''
        e= edge_data
        l,r,b,t = [e[1],e[1]+e[3],e[2],e[2]+e[4]] # get info to convert this to trace cell type
        tc = TraceCell(left=l, right=r, bottom=b, top=t)
        tc.z = elv
        tc.thick = dz * 1000
        tc.name = e[5]
        if self.ori_map[tc.name] == 'H':
            # LEFT AND RIGHT NODE
            l_locs = (l,b+e[4]/2,elv)
            r_locs = (r,b+e[4]/2,elv)
            self.add_and_update_nodes(l_locs,isl_name,tc,type = 'trace')
            self.add_and_update_nodes(r_locs,isl_name,tc,type = 'trace')
        elif self.ori_map[tc.name] == 'V':
            # TOP AND BOTTOM NODES
            b_locs = (l+e[3]/2,b,elv)
            t_locs = (l+e[3]/2,t,elv)
            self.add_and_update_nodes(t_locs,isl_name,tc,type = 'trace')
            self.add_and_update_nodes(b_locs,isl_name,tc,type = 'trace')
        return tc
    
    def handle_orthogonal_trace(self,t1,t2,o1,isl_name,elv):
        '''
        This function create a new node (*) at the center of the corner rectangle forming by 2 traces
        Then the initial nodes generated by these traces are removed (to maximize loop length)

        -----------------------------------|
        X                              *   X 
        -------------------------------X---|
                                    |      |
                                    |      |
                                    |      |
                                    |--X---|
        * : added node
        X : to be removed
        '''
        if o1 == 'H': # o2 =='V'
            c_loc = (t2.left+t2.width/2,t1.bottom+t1.height/2,elv)
            self.graph.add_node(self.node_id,locs= c_loc,island = isl_name, type = 'trace',parent = [t1,t2])
            self.tbl_isl_loc_net[isl_name][c_loc] = self.node_id
            self.node_id +=1
            # form a list of initial nodes (from naive graph nodes generation)
            p1 = (t1.left,c_loc[1],elv)
            p2 = (t1.right,c_loc[1],elv)
            p3 = (c_loc[0],t2.top,elv)
            p4 = (c_loc[0],t2.bottom,elv)
            rm_list = [p1,p2,p3,p4]
            for rm_pt in rm_list:
                if rm_pt in self.tbl_isl_loc_net[isl_name]:# if this node is in this graph
                    rm_node = self.tbl_isl_loc_net[isl_name][rm_pt]
                    self.graph.remove_node(rm_node)
                    del self.tbl_isl_loc_net[isl_name][rm_pt] # remove from table too
        else:
            c_loc = (t1.left+t1.width/2,t2.bottom+t2.height/2,elv)
            self.graph.add_node(self.node_id,locs=c_loc,island = isl_name, type = 'trace',parent = [t1,t2])
            self.tbl_isl_loc_net[isl_name][c_loc] = self.node_id
            self.node_id +=1
            p1 = (t2.left,c_loc[1],elv)
            p2 = (t2.right,c_loc[1],elv)
            p3 = (c_loc[0],t1.top,elv)
            p4 = (c_loc[0],t1.bottom,elv)
            rm_list = [p1,p2,p3,p4]
            for rm_pt in rm_list:
                if rm_pt in self.tbl_isl_loc_net[isl_name]:# if this node is in this graph
                    rm_node = self.tbl_isl_loc_net[isl_name][rm_pt]
                    self.graph.remove_node(rm_node)
                    del self.tbl_isl_loc_net[isl_name][rm_pt] # remove from table too



    
    def form_wire_frame(self, island = None,isl_name= '',layer_name='',elv=0,dz=0):
        '''
        The objective is to form a simplest wireframe graph, this is used to get the inital orientation of all traces
        island: layout info object
        isl_name: name of the current isl
        elv: elevation levels of elements on top and bottom of this trace
        form a simple wire frame from the iniital layout
        '''
        #print("elevation:",elv , "isl name", isl_name)
        
        elements = island.elements
        el_names = island.element_names
        debug = False
        pairs = {}
        # Find the longest extension for the trace based on the layout infomation
        # Define initial orientation
        traces = []
        mnodes_mode = 0 # mannual mode: 0 CS mode:1
        self.tbl_isl_loc_net[isl_name] = {}
        cs_node_map = {} # each table will have an island name, node Id --> node info (same for all meshes)
        for e in elements: # get all traces in this island
            if self.find_ori:
                self.find_trace_ori(e)
            '''
            l,r,b,t = [e[1],e[1]+e[3],e[2],e[2]+e[4]] # get info to convert this to trace cell type
            tc = TraceCell(left=l, right=r, bottom=b, top=t)
            tc.z = elv
            tc.thick = dz * 1000
            tc.name = e[5]
            '''
            if mnodes_mode == 0: # Mannually adding the nodes based on the initial layout
                tc = self.create_strace(edge_data= e, isl_name= isl_name, elv= elv, dz = dz)
                '''  # HANDLE PLANAR TRACE LATER HERE
                else: # planar trace
                    self.graph.add_node(self.node_id, locs = [l,b,elv],island = isl_name, type = 'trace',parent = [tc])
                    isl_nodes.append(self.node_id)
                    self.node_id +=1
                    self.graph.add_node(self.node_id, locs = [l,t,elv],island = isl_name, type = 'trace',parent = [tc])
                    isl_nodes.append(self.node_id)
                    self.node_id +=1
                    self.graph.add_node(self.node_id, locs = [r,b,elv],island = isl_name, type = 'trace',parent = [tc])
                    isl_nodes.append(self.node_id)
                    self.node_id +=1
                    self.graph.add_node(self.node_id, locs = [r,t,elv],island = isl_name, type = 'trace',parent = [tc])
                    isl_nodes.append(self.node_id)
                    self.node_id +=1
                '''
            traces.append(tc)
        # This is the first step to form the mesh table need to maintain this for every layout
        if mnodes_mode == 1: 
            for n in island.mesh_nodes:
                locs = [n.pos[0],n.pos[1],elv]
                self.graph.add_node(self.node_id,locs = locs)
                self.node_id+=1
        
        # Basic connections using PS1 idea # TODO: update for generic method
        num_t = len(traces) # number of traces
        for i in range(num_t):
            t1 = traces[i]
            for k in range(i+1,num_t,1):
                t2 = traces[k]
                o1 = self.ori_map[t1.name]
                o2 = self.ori_map[t2.name]
                if o1!= 'P' and o2!='P':
                    if o1 == o2: 
                        continue 
                    else: # Handle orthogonal traces ( L CONNECTION)
                        #add corner anchor to the graph
                        self.handle_orthogonal_trace(t1,t2,o1,isl_name,elv)
                        

        
        
        
        # Find all nets on this island
        for sh in self.hier.sheets:
            group = sh.parent.parent  # Define the trace island (containing a sheet)
            sheet_data = sh.data
            if isl_name == group.name:  # means if this sheet is in this island
                if not (group in self.comp_nodes):  # Create a list in dictionary to store all hierarchy node for each group
                    self.comp_nodes[group] = []

                comp = sh.data.component  # Get the component data from sheet.
                # print "C_DICT",len(self.comp_dict),self.comp_dict
                if comp != None and not (comp in self.comp_dict):  # In case this is an component with multiple pins
                    if not (sheet_data.net in self.comp_net_id):
                        comp.build_graph(mode =1) # Evaluate bondwires edge data here... 
                        # Get x,y,z positions
                        x, y = sheet_data.rect.center()
                        z = sheet_data.z
                        for tc in traces:
                            if tc.encloses(x , y) and z == (tc.z + tc.thick):
                                #print("ADD a component node",sheet_data.net,tc.z)
                                self.add_node_tracecell(tc,x,y,isl_name) # merge the z level to trace for simplification in connection
                                        
                        self.comp_net_id[sheet_data.net] = self.node_id-1
                        self.comp_dict[comp] = 1
                    for n in comp.net_graph.nodes(data=True):  # node without parents
                        sheet_data = n[1]['node']
                        if sheet_data.node == None:  # floating net
                            x, y = sheet_data.rect.center()
                            z = sheet_data.z
                            # print "CP",cp
                            if not (sheet_data.net in self.comp_net_id):
                                # print self.node_count
                                # floating node
                                #print ('float',sheet_data.net)
                                self.add_node_float(loc = (x,y,z), net = sheet_data.net,isl_name= isl_name)
                                self.comp_dict[comp] = 1
                else:  # In case this is simply a lead connection
                    if not (sheet_data.net in self.comp_net_id):

                        # Get x,y,z positions
                        x, y = sheet_data.rect.center()
                        z = sheet_data.z
                        for tc in traces:  # find which trace cell includes this component
                            #print('tc-z',tc.z)
                            #print(sheet_data.net, z)
                            if tc.encloses(x , y ):
                                #self.graph.add_node(self.node_id,locs=[x,y,tc.z],island = isl_name)
                                self.add_node_tracecell(tc,x,y,isl_name)  # merge the z level to trace for simplification in connection
                        self.comp_net_id[sheet_data.net] = self.node_id-1
                        #self.comp_nodes[group].append(cp_node)
                        #self.comp_dict[comp] = 1

                    
        
        for k in self.tbl_isl_loc_net[isl_name]:
            n = self.tbl_isl_loc_net[isl_name][k]
            self.tbl_isl_node_connection[n] = {'N':-1,'S':-1,'E':-1,'W':-1}
        for k1 in self.tbl_isl_loc_net[isl_name]: 
            n1 =  self.tbl_isl_loc_net[isl_name][k1]
            dN = 1e6
            dS = 1e6
            dE = 1e6
            dW = 1e6
            remE = -1
            remW = -1
            remN = -1
            remS = -1
            dir = None
            for k2 in self.tbl_isl_loc_net[isl_name]:
                n2 =  self.tbl_isl_loc_net[isl_name][k2]
                if n1 != n2 and not(self.graph.has_edge(n2,n1)): # find the best node to connect
                    p1 = self.graph.nodes[n1]['locs']
                    p2 = self.graph.nodes[n2]['locs']
                    t1 = self.graph.nodes[n1]['type']
                    t2 = self.graph.nodes[n1]['type']
                    pr1 = self.graph.nodes[n1]['parent']
                    pr2 = self.graph.nodes[n2]['parent']
                   
                    
                        
                    if t1=='hier' or t2 == 'hier' or pr1 == None or pr2==None: # hierachical node will be handle later
                        continue    
                    
                    # if they are both trace's node
                    set1 = set(pr1)
                    intersect = set1.intersection(pr2)
                    if len(intersect) == 0: # no intersection between parent traces, means no direct connection
                        #print("no connect",n1,n2)
                        continue
                    
                    if p1[2]!= p2[2]: # not same level
                        continue
                    else:
                        if p1[0] == p2[0]: # horizontal connection
                            dis_ = abs(p1[1]-p2[1])
                            if p2[1] > p1[1]: # check North connection
                                if self.tbl_isl_node_connection[n1]['N'] == -1:       
                                    dir= 'N'      
                                else:
                                    continue
                            else: # check South connection
                                if self.tbl_isl_node_connection[n1]['S'] == -1:      
                                    dir= 'S'      
                                else:
                                    continue
                            
                        elif p1[1]==p2[1]:
                            dir = 'V'
                            dis_ = abs(p1[0]-p2[0])
                            if p2[0] > p1[0]: # check East connection
                                if self.tbl_isl_node_connection[n1]['E'] == -1 :      
                                    dir= 'E'      
                                else:
                                    continue
                            else: # check West connection
                                dis_ = abs(p1[0]-p2[0])
                                if self.tbl_isl_node_connection[n1]['W'] == -1 :     
                                    dir= 'W'      
                                else:
                                    continue
                        else:
                            continue
                        if dir!= None:
                            if dir == 'E' and dis_ < dE:
                                dE = dis_
                                remE = n2
                            elif dir == 'W' and dis_ < dW:
                                dW = dis_
                                remW = n2
                            elif dir == 'N' and dis_ < dN:
                                dN = dis_
                                remN = n2
                            elif dir=='S'and dis_ < dS:
                                dS = dis_
                                remS = n2
                            
                else:
                    continue
            rem_list = [remN,remS,remE,remW]
            for i in range(4):
                r = rem_list[i]
                if r!=-1:
                    if i == 0 : #  North
                        self.tbl_isl_node_connection[n1]['N']=r 
                        self.tbl_isl_node_connection[r]['S']=n1 
                    elif i == 1 : #  South
                        self.tbl_isl_node_connection[n1]['S']=r
                        self.tbl_isl_node_connection[r]['N']=n1
                    elif i == 2 : #  East
                        self.tbl_isl_node_connection[n1]['E']=r
                        self.tbl_isl_node_connection[r]['W']=n1
                    elif i == 3 : #  West
                        self.tbl_isl_node_connection[n1]['W']=r
                        self.tbl_isl_node_connection[r]['E']=n1
                    e_data = EdgeData()
                    p_list1 = self.graph.nodes[n1]['parent']
                    p_list2 = self.graph.nodes[r]['parent']
                    set1 = set(p_list1)
                    set2 = set(p_list2)
                    if set1 & set2:
                        ptrace = set1 & set2
                        ptrace = list(ptrace)[0]
                    else:
                        ptrace = None
                    
                    #print(n1,r)
                    #print("parent_trace",ptrace.name)
                    #print(ptrace.eval_length())
                    self.graph.add_edge(n1,r,e_type = 'trace',p_trace= ptrace,res_int=1)
                
    def plot(self,option = "all", mode = 1, isl= "",pos = {},graph = None,save=False,mem_file=''):
        '''In debug mode, a mem_file iostream will be used to store graph information and then save to word document'''
        if mode == 1: # save figure as file   
            plt.figure("digraph")
            nx.draw(self.digraph,self.pos,with_labels=True)
            
        elif mode == 2: #save figure as data to replot
            fig1 = plt.figure(isl)
            ax1 = fig1.add_subplot(111)
            ax1.set_title(isl)
            nx.draw(graph,pos,with_labels=True) # mesh graph
            
            

        elif mode == 3:
            # plot by each elevation level.
            for ele in self.ele_lst:
                # for each elevation make a copy of digraph and self.pos
                n_pos = deepcopy(self.pos)
                n_graph = deepcopy(self.digraph)
                for n in self.digraph.nodes: # Only plot the nodes on same level and save to pickle
                    z =  self.graph.nodes[n]['locs'][2]
                    if z != ele:
                        del n_pos[n]
                        n_graph.remove_node(n)
                fig1 = plt.figure(ele)
                ax1 = fig1.add_subplot(111)
                ax1.set_title("digraph on ele:" + str(ele))
                nx.draw(n_graph,n_pos,with_labels=True) # mesh graph
        
        # FOR 2D layout netlist debugging only
        
        elif mode ==4:
            fig = plt.figure("net_graph")
            ax1 = fig.add_subplot(111)
            ax1.set_title('NET GRAPH')
            nx.draw_networkx(self.net_graph, pos=self.net_2d_pos,
                             with_labels=True, node_size=50, font_size=12)
        elif mode == 5:
            fig = plt.figure(1,dpi=200)
            ax1 = fig.add_subplot(111)
            ax1.set_title("NETLIST")
            # FIRST WE SORT THE XS AND YS VALUES FOR EACH BUNDLE
            xs_to_xp = {}
            ys_to_yp = {}
            mult = 5# a multiplier to increase the edge distance
            xp = []
            yp = []
            new_2d_pos_scaled= {}
            net_labels = {}
            for p in self.net_2d_pos:
                point = self.net_2d_pos[p]
                xp.append(point[0])
                yp.append(point[1])
            xp = list(set(xp))
            xp.sort()
            yp = list(set(yp))
            yp.sort()
            for ix in range(len(xp)):
                xs_to_xp[xp[ix]] = ix*mult
            for iy in range(len(yp)):
                ys_to_yp[yp[iy]] = iy*mult
            # upodate the postion 
            for p in self.net_2d_pos:
                point = self.net_2d_pos[p]
                point_scaled = (xs_to_xp[point[0]],ys_to_yp[point[1]])
                new_2d_pos_scaled[p] = point_scaled 
            print(new_2d_pos_scaled)
            for e in self.net_graph.edges(data = True):
                edata = e[2]['data']
                eval = e[2]
                if 'fw' in edata['type']:
                    R = round(eval['res']*1e3,2)
                    L = round(eval['ind']*1e9,2)
                    line = "{}m-{}n".format(R,L)     
                    net_labels[(e[0],e[1])] = line           
            nx.draw_networkx(self.net_graph, pos=new_2d_pos_scaled,
                             with_labels=True, node_size=10, font_size=10)
            nx.draw_networkx_edge_labels(self.net_graph,pos = new_2d_pos_scaled,edge_labels=net_labels,
                                         font_size=8,font_color='black',font_family='arial')
            plt.autoscale()  
        if save == True:
            plt.savefig(mem_file,format='png')
        else:
            plt.show()
    def save_plot_pickle(self,ax,fname):
        pickle.dump(ax,open(fname,"wb"))

    
    
    def find_all_paths(self,src,sink):
        '''
        Once the wire frame is formed, find all loops to define directions of all edges in the Digraph
        '''
        s1 = self.comp_net_id[src]
        s2 = self.comp_net_id[sink]
        #print (nx.has_path(self.graph,s1,s2))
        #print(self.graph.has_node(0))
        paths = nx.all_simple_paths(self.graph,s1,s2)
        for p in paths:
            #print(p)
            for i in range(len(p)-1):
                if not(self.digraph.has_edge(p[i],p[i+1])):
                    self.digraph.add_edge(p[i],p[i+1])
        
        # update the graph by combining all contracted node pairs
        # the 2nd node v will be merged into the first node u
        
        for k in self.contracted_nodes:
            self.digraph = nx.contracted_nodes(self.digraph,k,self.contracted_nodes[k],self_loops = False)
        if self.debug:
            memfile = io.BytesIO()
            self.plot(mode=3,save=True,mem_file=memfile)
            self.doc_handle_figure(memfile=memfile,fig_heading="Initial digraph")
        
        '''
        # modify to create more edges in x and y 
        # first get all x and y locs and sort them:
        x_locs = []
        y_locs = []
        for pos in self.pos.values():
            x_locs.append(pos[0])
            y_locs.append(pos[1])
        x_locs = list(set(x_locs))
        x_locs.sort()
        y_locs = list(set(y_locs))
        y_locs.sort()

        for e in self.digraph.edges(data=True): # Loop through all edges, only split trace edges 
            # get edge data from graph
            edata = self.graph.get_edge_data(e[0],e[1])
            locx = [self.pos[e[0]][0],self.pos[e[1]][0]]
            locy = [self.pos[e[0]][1],self.pos[e[1]][1]]
            if locy[0] == locy[1]:# horizontal edge
                left = min(locx)
                right = max(locx)
                
                for i in range(len(x_locs)):
                    if x_locs[i] < right and x_locs[i]>=left:
                        
        '''                


    def find_trace_ori(self,el_data):
        # predefine layout orientation based on width and height info, later will update one more time through current direction
        if el_data[3] >= 2 *el_data[4]:
            self.ori_map[el_data[5]] = 'H'
        elif el_data[4] >= 2 *el_data[3]:
            self.ori_map[el_data[5]] = 'V'
        else:
            self.ori_map[el_data[5]] = 'P'

def solve_loop_models_parallel(all_loops=None):
    num = int(cpu_count()/2)
    with Pool(num) as p:
        results = p.map(solve_single_loop,all_loops)
    return results
def solve_single_loop(loop_model):
    loop_model.form_mesh_matrix()
    loop_model.update_P(1e9)
    loop_model.solve_linear_systems()
    return loop_model
def load_pickle_plot(fname):
    ax1 = pickle.load(open("mesh.p","rb"))
    ax2 = pickle.load(open("digraph.p","rb"))
    
    plt.show()
if __name__ == '__main__':
    l1 = LoopDefinition()
    l1.test_struct_L_shape()  
